import os
import requests
from datetime import datetime
from docx import Document

# 1. 환경 변수 로드
GEMINI_KEY = os.getenv('GEMINI_API_KEY')
TELE_TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')
CHAT_ID = os.getenv('TELEGRAM_CHAT_ID')

def get_working_model():
    """대표님 계정에서 현재 작동 가능한 모델명을 찾아냅니다."""
    list_url = f"https://generativelanguage.googleapis.com/v1beta/models?key={GEMINI_KEY}"
    try:
        res = requests.get(list_url)
        models_data = res.json()
        # 'generateContent' 기능을 지원하는 모델만 필터링
        valid_models = [
            m['name'].split('/')[-1] for m in models_data.get('models', []) 
            if 'generateContent' in m.get('supportedGenerationMethods', [])
        ]
        
        # 선호도 순으로 매칭
        for preferred in ["gemini-1.5-flash-latest", "gemini-1.5-flash", "gemini-1.5-pro", "gemini-pro"]:
            if preferred in valid_models:
                return preferred
        return valid_models[0] if valid_models else None
    except:
        return "gemini-1.5-flash-latest" # 실패 시 기본값 설정

def get_ai_data(model_name):
    keywords = ["외식업 AI 자동화", "심리상담 AI 트렌드", "비즈니스 에이전트", "외식업 AI활용", "AI도입"]
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={GEMINI_KEY}"
    headers = {'Content-Type': 'application/json'}
    reports = []
    
    print(f"🚀 확정된 모델 [{model_name}]으로 수집을 시작합니다.")
    
    for kw in keywords:
        prompt = f"30년 경영자 관점에서 '{kw}' 관련 최신 트렌드 1개를 제목, 3줄 요약, 경영 인사이트 순으로 정리해줘."
        data = {"contents": [{"parts": [{"text": prompt}]}]}
        try:
            res = requests.post(url, headers=headers, json=data, timeout=60)
            res.raise_for_status()
            content = res.json()['candidates'][0]['content']['parts'][0]['text']
            reports.append(content)
            print(f"✅ '{kw}' 분석 완료")
        except Exception as e:
            print(f"❌ '{kw}' 건너뜀: {e}")
    return reports

def create_and_send(results):
    # 1. Word 파일 생성
    file_name = f"AI_Briefing_{datetime.now().strftime('%m%d')}.docx"
    doc = Document()
    doc.add_heading('최장규 대표의 AI 비즈니스 인텔리전스', 0)
    for r in results:
        doc.add_paragraph(r)
        doc.add_paragraph("-" * 50)
    doc.save(file_name)
    
    # 2. 텔레그램 전송
    send_url = f"https://api.telegram.org/bot{TELE_TOKEN}/sendDocument"
    with open(file_name, 'rb') as f:
        payload = {'chat_id': CHAT_ID, 'caption': f"📅 {datetime.now().strftime('%Y-%m-%d')} 리포트 도착"}
        files = {'document': f}
        response = requests.post(send_url, data=payload, files=files)
        if response.status_code == 200:
            print("✨ 텔레그램 전송 성공!")
        else:
            print(f"❌ 전송 실패: {response.text}")

if __name__ == "__main__":
    if all([GEMINI_KEY, TELE_TOKEN, CHAT_ID]):
        target = get_working_model()
        if target:
            data = get_ai_data(target)
            if data:
                create_and_send(data)
        else:
            print("❌ 사용 가능한 AI 모델이 없습니다.")
    else:
        print("❌ 설정(Secrets) 확인 필요")
